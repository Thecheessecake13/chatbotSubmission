from pathlib import Path

import fitz
from docx import Document as DocxDocument

SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}


class ExtractionError(Exception):
    pass


def extract_text(path: Path) -> list[tuple[str, int | None]]:
    suffix = path.suffix.lower()
    if suffix == '.pdf':
        return _extract_pdf(path)
    if suffix == '.docx':
        return _extract_docx(path)
    raise ExtractionError(f'Unsupported file extension: {suffix}')


def _extract_pdf(path: Path) -> list[tuple[str, int | None]]:
    try:
        doc = fitz.open(path)
    except Exception as exc:
        raise ExtractionError('Could not open PDF. The file may be encrypted or corrupt.') from exc

    pages: list[tuple[str, int | None]] = []
    try:
        for index, page in enumerate(doc, start=1):
            text = page.get_text('text').strip()
            if text:
                pages.append((text, index))
    finally:
        doc.close()

    if not pages:
        raise ExtractionError('No searchable text found. Scanned PDFs need OCR, which is intentionally not enabled in this assignment build.')
    return pages


def _extract_docx(path: Path) -> list[tuple[str, int | None]]:
    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        raise ExtractionError('Could not open DOCX. The file may be corrupt.') from exc

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(' | '.join(cells))

    full_text = '\n'.join(parts).strip()
    if not full_text:
        raise ExtractionError('No text found in DOCX.')
    return [(full_text, None)]
